"""File parsers for different document formats."""

from __future__ import annotations

import csv
import json
from abc import ABC, abstractmethod
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from cgc.core.schema import DataType


@dataclass
class ParsedContent:
    """Result of parsing a file."""

    text: str  # Extracted text content
    rows: list[dict[str, Any]] | None = None  # Structured data if applicable
    fields: list[tuple[str, DataType]] | None = None  # Inferred schema
    metadata: dict[str, Any] | None = None  # File-specific metadata
    page_count: int | None = None  # For PDFs
    sheet_names: list[str] | None = None  # For Excel


class FileParser(ABC):
    """Base class for file parsers."""

    @abstractmethod
    def parse(self, content: bytes, filename: str) -> ParsedContent:
        """Parse file content and return extracted data."""
        ...

    @abstractmethod
    def supports(self, extension: str) -> bool:
        """Check if this parser supports the file extension."""
        ...


class TextParser(FileParser):
    """Parser for plain text files."""

    EXTENSIONS = {".txt", ".log", ".text"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text = content.decode("utf-8", errors="replace")
        lines = text.split("\n")
        return ParsedContent(
            text=text,
            fields=[("content", DataType.TEXT), ("line_count", DataType.INTEGER)],
            metadata={"line_count": len(lines), "char_count": len(text)},
        )


class MarkdownParser(FileParser):
    """Parser for Markdown files."""

    EXTENSIONS = {".md", ".markdown", ".mdown"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text = content.decode("utf-8", errors="replace")
        lines = text.split("\n")

        # Extract headers for structure
        headers = [line for line in lines if line.startswith("#")]

        return ParsedContent(
            text=text,
            fields=[("content", DataType.TEXT), ("headers", DataType.ARRAY)],
            metadata={
                "line_count": len(lines),
                "char_count": len(text),
                "headers": headers[:10],  # First 10 headers
            },
        )


class CsvParser(FileParser):
    """Parser for CSV/TSV files."""

    EXTENSIONS = {".csv", ".tsv"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text = content.decode("utf-8", errors="replace")

        # Detect delimiter
        delimiter = "\t" if filename.endswith(".tsv") else ","

        reader = csv.DictReader(StringIO(text), delimiter=delimiter)
        rows = list(reader)

        # Infer field types from first 100 rows
        fields = []
        if rows:
            for col in rows[0].keys():
                dtype = self._infer_type([r.get(col) for r in rows[:100]])
                fields.append((col, dtype))

        return ParsedContent(
            text=text,
            rows=rows,
            fields=fields,
            metadata={"row_count": len(rows), "column_count": len(fields)},
        )

    def _infer_type(self, values: list) -> DataType:
        """Infer data type from sample values."""
        for val in values:
            if val is None or val == "":
                continue

            # Try integer
            try:
                int(val)
                return DataType.INTEGER
            except (ValueError, TypeError):
                pass

            # Try float
            try:
                float(val)
                return DataType.FLOAT
            except (ValueError, TypeError):
                pass

            # Check boolean
            if str(val).lower() in ("true", "false"):
                return DataType.BOOLEAN

        return DataType.STRING


class JsonParser(FileParser):
    """Parser for JSON files."""

    EXTENSIONS = {".json"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text = content.decode("utf-8", errors="replace")
        data = json.loads(text)

        rows = None
        fields = None

        if isinstance(data, list) and data and isinstance(data[0], dict):
            rows = data
            fields = [(k, self._infer_type(v)) for k, v in data[0].items()]
        elif isinstance(data, dict):
            rows = [data]
            fields = [(k, self._infer_type(v)) for k, v in data.items()]

        return ParsedContent(
            text=text,
            rows=rows,
            fields=fields,
            metadata={"type": "array" if isinstance(data, list) else "object"},
        )

    def _infer_type(self, value: Any) -> DataType:
        if value is None:
            return DataType.STRING
        if isinstance(value, bool):
            return DataType.BOOLEAN
        if isinstance(value, int):
            return DataType.INTEGER
        if isinstance(value, float):
            return DataType.FLOAT
        if isinstance(value, list):
            return DataType.ARRAY
        if isinstance(value, dict):
            return DataType.JSON
        return DataType.STRING


class JsonlParser(FileParser):
    """Parser for JSONL/NDJSON files."""

    EXTENSIONS = {".jsonl", ".ndjson"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text = content.decode("utf-8", errors="replace")
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        rows = [json.loads(line) for line in lines]

        fields = None
        if rows and isinstance(rows[0], dict):
            fields = [(k, DataType.JSON) for k in rows[0].keys()]

        return ParsedContent(
            text=text,
            rows=rows,
            fields=fields,
            metadata={"row_count": len(rows)},
        )


class PdfParser(FileParser):
    """Parser for PDF files."""

    EXTENSIONS = {".pdf"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        try:
            import pypdf
        except ImportError:
            return ParsedContent(
                text="[PDF parsing requires pypdf: pip install pypdf]",
                metadata={"error": "pypdf not installed"},
            )

        try:
            reader = pypdf.PdfReader(BytesIO(content))
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")

            text = "\n\n---PAGE BREAK---\n\n".join(pages)

            return ParsedContent(
                text=text,
                fields=[("content", DataType.TEXT), ("page", DataType.INTEGER)],
                metadata={
                    "page_count": len(reader.pages),
                    "title": reader.metadata.title if reader.metadata else None,
                    "author": reader.metadata.author if reader.metadata else None,
                },
                page_count=len(reader.pages),
            )
        except Exception as e:
            return ParsedContent(
                text=f"[PDF parsing error: {e}]",
                metadata={"error": str(e)},
            )


class DocxParser(FileParser):
    """Parser for DOCX files."""

    EXTENSIONS = {".docx"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        try:
            from docx import Document
        except ImportError:
            return ParsedContent(
                text="[DOCX parsing requires python-docx: pip install python-docx]",
                metadata={"error": "python-docx not installed"},
            )

        try:
            doc = Document(BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)

            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    table_rows.append([cell.text for cell in row.cells])
                tables_text.append(table_rows)

            return ParsedContent(
                text=text,
                fields=[("content", DataType.TEXT)],
                metadata={
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                    "tables": tables_text[:5],  # First 5 tables
                },
            )
        except Exception as e:
            return ParsedContent(
                text=f"[DOCX parsing error: {e}]",
                metadata={"error": str(e)},
            )


class ExcelParser(FileParser):
    """Parser for XLS/XLSX files."""

    EXTENSIONS = {".xls", ".xlsx"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        try:
            import openpyxl
        except ImportError:
            return ParsedContent(
                text="[Excel parsing requires openpyxl: pip install openpyxl]",
                metadata={"error": "openpyxl not installed"},
            )

        try:
            wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
            sheet_names = wb.sheetnames

            all_rows = []
            all_text = []

            for sheet_name in sheet_names:
                sheet = wb[sheet_name]
                all_text.append(f"=== Sheet: {sheet_name} ===")

                # Get headers from first row
                headers = []
                for cell in sheet[1]:
                    headers.append(str(cell.value) if cell.value else f"col_{cell.column}")

                # Get data rows
                for row in sheet.iter_rows(min_row=2, values_only=True):
                    row_dict = dict(zip(headers, row))
                    row_dict["_sheet"] = sheet_name
                    all_rows.append(row_dict)
                    all_text.append("\t".join(str(v) if v else "" for v in row))

            text = "\n".join(all_text)

            # Infer fields from first sheet
            fields = None
            if all_rows:
                first_row = all_rows[0]
                fields = [(k, DataType.STRING) for k in first_row.keys()]

            return ParsedContent(
                text=text,
                rows=all_rows,
                fields=fields,
                metadata={
                    "sheet_count": len(sheet_names),
                    "total_rows": len(all_rows),
                },
                sheet_names=sheet_names,
            )
        except Exception as e:
            return ParsedContent(
                text=f"[Excel parsing error: {e}]",
                metadata={"error": str(e)},
            )


class CodeParser(FileParser):
    """Parser for code files (treats as text with metadata)."""

    EXTENSIONS = {
        ".py", ".js", ".ts", ".jsx", ".tsx",
        ".java", ".go", ".rs", ".c", ".cpp", ".h",
        ".rb", ".php", ".swift", ".kt", ".scala",
        ".html", ".css", ".scss", ".sql", ".sh",
        ".yaml", ".yml", ".toml", ".ini", ".xml",
    }

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        text = content.decode("utf-8", errors="replace")
        lines = text.split("\n")
        ext = Path(filename).suffix.lower()

        # Count comments (simple heuristic)
        comment_chars = {
            ".py": "#",
            ".js": "//", ".ts": "//", ".jsx": "//", ".tsx": "//",
            ".java": "//", ".go": "//", ".rs": "//",
            ".c": "//", ".cpp": "//", ".h": "//",
            ".rb": "#", ".php": "//", ".sh": "#",
        }
        comment_prefix = comment_chars.get(ext, "#")
        comment_lines = sum(1 for line in lines if line.strip().startswith(comment_prefix))

        return ParsedContent(
            text=text,
            fields=[("content", DataType.TEXT)],
            metadata={
                "language": ext[1:],  # Remove dot
                "line_count": len(lines),
                "comment_lines": comment_lines,
                "char_count": len(text),
            },
        )


class ParquetParser(FileParser):
    """Parser for Parquet files."""

    EXTENSIONS = {".parquet"}

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.EXTENSIONS

    def parse(self, content: bytes, filename: str) -> ParsedContent:
        try:
            import pyarrow.parquet as pq
        except ImportError:
            return ParsedContent(
                text="[Parquet parsing requires pyarrow: pip install pyarrow]",
                metadata={"error": "pyarrow not installed"},
            )

        try:
            table = pq.read_table(BytesIO(content))
            df = table.to_pandas()
            rows = df.to_dict(orient="records")

            fields = []
            for col in df.columns:
                dtype = DataType.STRING
                if df[col].dtype in ("int64", "int32"):
                    dtype = DataType.INTEGER
                elif df[col].dtype in ("float64", "float32"):
                    dtype = DataType.FLOAT
                elif df[col].dtype == "bool":
                    dtype = DataType.BOOLEAN
                fields.append((col, dtype))

            text = df.head(100).to_string()

            return ParsedContent(
                text=text,
                rows=rows,
                fields=fields,
                metadata={
                    "row_count": len(rows),
                    "column_count": len(df.columns),
                },
            )
        except Exception as e:
            return ParsedContent(
                text=f"[Parquet parsing error: {e}]",
                metadata={"error": str(e)},
            )


# Registry of all parsers
PARSERS: list[FileParser] = [
    TextParser(),
    MarkdownParser(),
    CsvParser(),
    JsonParser(),
    JsonlParser(),
    PdfParser(),
    DocxParser(),
    ExcelParser(),
    CodeParser(),
    ParquetParser(),
]


def get_parser(filename: str) -> FileParser | None:
    """Get appropriate parser for a file."""
    ext = Path(filename).suffix.lower()
    for parser in PARSERS:
        if parser.supports(ext):
            return parser
    return None


def parse_file(content: bytes, filename: str) -> ParsedContent:
    """Parse a file and return extracted content."""
    parser = get_parser(filename)
    if parser:
        return parser.parse(content, filename)

    # Fallback: try to decode as text
    try:
        text = content.decode("utf-8", errors="replace")
        return ParsedContent(
            text=text,
            metadata={"type": "unknown", "size": len(content)},
        )
    except Exception:
        return ParsedContent(
            text=f"[Binary file: {len(content)} bytes]",
            metadata={"type": "binary", "size": len(content)},
        )
